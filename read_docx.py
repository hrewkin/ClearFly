from docx import Document
import os

path = r'c:\Users\dima\Desktop\Смирнов матьего\Требования_АС_Чистое_Небо_ласт_апдейт_на_паре.docx'
doc = Document(path)

out_path = r'c:\Users\dima\Desktop\Смирнов матьего\requirements_text.txt'

with open(out_path, 'w', encoding='utf-8') as f:
    f.write("=== PARAGRAPHS ===\n")
    for p in doc.paragraphs:
        if p.text.strip():
            f.write(p.text + "\n")
    
    f.write("\n=== TABLES ===\n")
    for i, table in enumerate(doc.tables):
        f.write(f"\n--- Table {i+1} ---\n")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            f.write(" | ".join(cells) + "\n")

print("Done! Output written to requirements_text.txt")
